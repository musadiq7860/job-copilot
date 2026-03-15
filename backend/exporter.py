from docx import Document
import io

def create_docx(bullets: str, cover_letter: str, candidate_name: str) -> bytes:
    doc = Document()

    doc.add_heading(f"{candidate_name} — Tailored Resume Bullets", level=1)
    doc.add_paragraph("")

    doc.add_heading("Rewritten Resume Bullets", level=2)
    for line in bullets.split("\n"):
        if line.strip():
            doc.add_paragraph(line.strip(), style="List Bullet")

    doc.add_paragraph("")
    doc.add_heading("Cover Letter", level=2)
    for para in cover_letter.split("\n"):
        if para.strip():
            doc.add_paragraph(para.strip())

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()